# FILE: intake/document_reader.py
from __future__ import annotations

from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
import os
import re


@dataclass
class DocChunk:
    doc_id: str
    chunk_id: int
    text: str
    meta: Dict[str, Any] = field(default_factory=dict)

    @property
    def evidence_id(self) -> str:
        return str(self.meta.get("evidence_id") or f"{self.doc_id}:chunk_{self.chunk_id}")

    @property
    def source_label(self) -> str:
        return str(self.meta.get("source_label") or "reading")


class DocumentReader:
    """
    Reads documents into clean text and splits them into structured chunks.

    Supported:
    - .txt, .md
    - .pdf (prefers PyMuPDF/fitz, falls back to pdfminer.six)
    - .docx (python-docx)

    Defensive:
    - Detects likely PDF-stream garbage and rejects it
    - Tries to preserve document structure
    - Adds richer metadata to chunks for later memory writing

    Convenience:
    - read_for_prompt() returns a compact package for benchmark runners
    - select_relevant_chunks() ranks chunks by lexical overlap with a query
    """

    _PDF_GARBAGE_MARKERS = (
        "%pdf", "endobj", "xref", "endstream", "stream", "/flatedecode", "/filter", "obj"
    )

    _CHAPTER_RE = re.compile(r"\bchapter\s+([ivxlcdm]+|\d+)\b", re.IGNORECASE)
    _SECTION_RE = re.compile(r"\b(section|lesson|part)\s+([ivxlcdm]+|\d+)\b", re.IGNORECASE)
    _HEADING_RE = re.compile(r"^(chapter|section|part)\b", re.IGNORECASE)
    _WORD_RE = re.compile(r"\S+")
    _TOKEN_RE = re.compile(r"[a-zA-Z0-9']+")

    # ------------------------
    # Raw readers
    # ------------------------
    def read_text_file(self, path: str) -> str:
        for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                with open(path, "r", encoding=encoding, errors="ignore") as f:
                    return f.read()
            except Exception:
                continue
        return ""

    def read_docx(self, path: str) -> str:
        try:
            from docx import Document  # python-docx
        except Exception as e:
            raise RuntimeError("python-docx not installed, cannot read .docx") from e

        doc = Document(path)
        parts: List[str] = []

        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        # Include table content too
        for table in doc.tables:
            for row in table.rows:
                row_cells = [(cell.text or "").strip() for cell in row.cells]
                row_text = " | ".join(cell for cell in row_cells if cell).strip()
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    def read_pdf(self, path: str) -> str:
        text = ""
        fitz_exc: Optional[Exception] = None

        # Prefer PyMuPDF
        try:
            import fitz  # type: ignore

            doc = fitz.open(path)
            parts: List[str] = []
            for page in doc:
                t = page.get_text("text") or ""
                t = t.strip()
                if t:
                    parts.append(t)
            text = "\n\n".join(parts)
        except Exception as e:
            fitz_exc = e
            text = ""

        if text:
            text = self._clean_text(text)
            if not self._looks_like_pdf_garbage(text) and self._is_readable_text(text):
                return text

        # Fallback: pdfminer.six
        try:
            from pdfminer.high_level import extract_text  # type: ignore

            text2 = extract_text(path) or ""
            text2 = self._clean_text(text2)
            if text2 and not self._looks_like_pdf_garbage(text2) and self._is_readable_text(text2):
                return text2
        except Exception:
            pass

        if fitz_exc:
            return ""
        return ""

    def read_document(self, path: str) -> str:
        path = self._clean_path(path)
        if not path:
            return ""
        if not os.path.exists(path):
            return ""

        _, ext = os.path.splitext(path)
        ext = ext.lower()

        try:
            if ext in (".txt", ".md", ".text"):
                return self._clean_text(self.read_text_file(path))
            if ext == ".docx":
                return self._clean_text(self.read_docx(path))
            if ext == ".pdf":
                return self.read_pdf(path)

            # Unknown extension: last-resort text read
            return self._clean_text(self.read_text_file(path))
        except Exception:
            return ""

    # ------------------------
    # Validation / cleanup
    # ------------------------
    def _clean_path(self, path: str) -> str:
        path = (path or "").strip()
        if len(path) >= 2 and path[0] == path[-1] and path[0] in {"'", '"'}:
            path = path[1:-1].strip()
        return path

    def _looks_like_pdf_garbage(self, text: str) -> bool:
        if not text:
            return True

        low = text.lower()

        hits = sum(1 for m in self._PDF_GARBAGE_MARKERS if m in low)
        if hits >= 2:
            return True

        printable = sum(1 for ch in text if 32 <= ord(ch) <= 126 or ch in "\n\r\t")
        ratio = printable / max(1, len(text))
        if ratio < 0.85:
            return True

        return False

    def _is_readable_text(self, text: str) -> bool:
        if not text:
            return False

        stripped = text.strip()
        if len(stripped) < 60:
            return False

        visible = sum(1 for ch in stripped if not ch.isspace())
        alpha = sum(1 for ch in stripped if ch.isalpha())
        digit = sum(1 for ch in stripped if ch.isdigit())

        if visible < 40:
            return False

        alpha_ratio = alpha / max(1, len(stripped))
        signal_ratio = (alpha + digit) / max(1, len(stripped))

        if alpha_ratio < 0.20 and signal_ratio < 0.28:
            return False

        return True

    def _clean_text(self, doc_text: str) -> str:
        if not doc_text:
            return ""

        doc_text = doc_text.replace("\x00", " ")
        doc_text = doc_text.replace("\r\n", "\n").replace("\r", "\n")
        doc_text = re.sub(r"[ \t]+", " ", doc_text)
        doc_text = re.sub(r"\n{3,}", "\n\n", doc_text)
        return doc_text.strip()

    # ------------------------
    # Metadata inference
    # ------------------------
    def _extract_chapter_hint(self, file_stem: str, text: str) -> Optional[str]:
        for src in (file_stem, text[:1200]):
            m = self._CHAPTER_RE.search(src or "")
            if m:
                return m.group(1)
        return None

    def _extract_section_hint(self, file_stem: str, text: str) -> Optional[str]:
        for src in (file_stem, text[:1200]):
            m = self._SECTION_RE.search(src or "")
            if m:
                return m.group(2)
        return None

    def _guess_title(self, file_stem: str, text: str) -> str:
        if file_stem and file_stem.strip():
            return file_stem.strip()

        for line in text.splitlines()[:8]:
            clean = line.strip()
            if 4 <= len(clean) <= 120:
                return clean

        return "Untitled Document"

    def _guess_book_title(self, file_stem: str, folder_path: str, text: str) -> Optional[str]:
        folder_name = os.path.basename(folder_path) if folder_path else ""
        candidates = [file_stem, folder_name]

        for c in candidates:
            if not c:
                continue
            clean = c.replace("_", " ").replace("-", " ").strip()
            if len(clean) >= 4 and "lesson" not in clean.lower():
                return clean

        for line in text.splitlines()[:6]:
            clean = line.strip()
            if 5 <= len(clean) <= 120 and "chapter" not in clean.lower():
                return clean

        return None

    def _infer_domain_from_text(self, text: str) -> str:
        tl = (text or "").lower()

        if re.search(r"\b(appearance|reality|philosophy|knowledge|doubt|truth|sense-data|perception)\b", tl):
            return "philosophy"
        if re.search(r"\b(flood|fire|earthquake|survival|danger|safety)\b", tl):
            return "survival"
        if re.search(r"\b(task|plan|organize|priority|schedule)\b", tl):
            return "productivity"
        if re.search(r"\b(biology|human|water|science|physics|chemistry)\b", tl):
            return "general"

        return "general"

    def _infer_file_type_label(self, path: str) -> str:
        _, ext = os.path.splitext(path or "")
        ext = ext.lower().strip()
        if ext == ".pdf":
            return "pdf"
        if ext == ".docx":
            return "docx"
        if ext in {".txt", ".md", ".text"}:
            return "text"
        return ext.lstrip(".") or "unknown"

    def _build_doc_id(self, path: str) -> str:
        file_name = os.path.basename(path or "").strip()
        if file_name:
            return file_name
        return "unknown_document"

    def _build_evidence_id(self, doc_id: str, chunk_id: int) -> str:
        return f"{doc_id}:chunk_{chunk_id}"

    def _build_chunk_label(self, file_name: str, chunk_id: int, total_chunks: int) -> str:
        return f"{file_name or 'document'} chunk {chunk_id + 1}/{max(1, total_chunks)}"

    def _extract_preview(self, text: str, max_chars: int = 180) -> str:
        clean = self._clean_text(text)
        if len(clean) <= max_chars:
            return clean
        return clean[: max_chars - 3].rstrip() + "..."

    def _infer_reader_name(self, path: str) -> str:
        _, ext = os.path.splitext(path or "")
        ext = ext.lower()
        if ext == ".pdf":
            return "document_reader.pdf"
        if ext == ".docx":
            return "document_reader.docx"
        return "document_reader.text"

    # ------------------------
    # Chunking
    # ------------------------
    def chunk_text(
        self,
        doc_text: str,
        *,
        words_per_chunk: int = 900,
        overlap_words: int = 120,
    ) -> List[str]:
        cleaned = self._clean_text(doc_text)
        if not cleaned:
            return []

        # Preserve paragraph structure for short documents
        paras = [p.strip() for p in cleaned.split("\n\n") if p.strip()]
        if len(cleaned) < 3000 and paras:
            return paras

        words = self._WORD_RE.findall(cleaned)
        if not words:
            return []

        words_per_chunk = max(80, int(words_per_chunk))
        overlap_words = max(0, min(int(overlap_words), words_per_chunk // 3))

        chunks: List[str] = []
        step = max(1, words_per_chunk - overlap_words)
        i = 0

        while i < len(words):
            chunk_words = words[i:i + words_per_chunk]
            chunk = " ".join(chunk_words).strip()
            if chunk:
                chunks.append(chunk)
            if i + words_per_chunk >= len(words):
                break
            i += step

        return chunks

    def _find_heading_hint(self, chunk_text: str) -> Optional[str]:
        lines = [ln.strip() for ln in chunk_text.splitlines() if ln and ln.strip()]
        for ln in lines[:5]:
            if len(ln) <= 100 and (self._HEADING_RE.search(ln) or ln.isupper()):
                return ln
        return None

    def _make_chunk_meta(
        self,
        *,
        path: str,
        doc_id: str,
        file_stem: str,
        chunk_id: int,
        chunk_text: str,
        total_chunks: int,
    ) -> Dict[str, Any]:
        folder_path = os.path.dirname(path) or ""
        file_name = os.path.basename(path)
        title = self._guess_title(file_stem, chunk_text)
        chapter_hint = self._extract_chapter_hint(file_stem, chunk_text)
        section_hint = self._extract_section_hint(file_stem, chunk_text)
        heading_hint = self._find_heading_hint(chunk_text)
        book_title = self._guess_book_title(file_stem, folder_path, chunk_text)
        domain_hint = self._infer_domain_from_text(chunk_text)
        file_type = self._infer_file_type_label(path)
        evidence_id = self._build_evidence_id(doc_id, chunk_id)
        chunk_label = self._build_chunk_label(file_name, chunk_id, total_chunks)

        reading_tags: List[str] = ["reading", "document"]
        if file_type:
            reading_tags.append(file_type)
        if chapter_hint:
            reading_tags.append(f"chapter_{chapter_hint}")
        if section_hint:
            reading_tags.append(f"section_{section_hint}")
        if domain_hint:
            reading_tags.append(domain_hint)

        dedup_tags = list(dict.fromkeys([t for t in reading_tags if t]))

        return {
            "path": path,
            "file_name": file_name,
            "file_stem": file_stem,
            "file_type": file_type,
            "folder_path": folder_path,
            "doc_id": doc_id,
            "title": title,
            "book_title_hint": book_title,
            "chapter_hint": chapter_hint,
            "section_hint": section_hint,
            "heading_hint": heading_hint,
            "domain_hint": domain_hint,
            "chunk_index": chunk_id,
            "total_chunks": total_chunks,
            "word_count": len(self._WORD_RE.findall(chunk_text)),
            "char_count": len(chunk_text),
            "source_label": "reading",
            "reader_name": self._infer_reader_name(path),
            "evidence_id": evidence_id,
            "chunk_label": chunk_label,
            "evidence_text": f"{file_name}:{chunk_id}",
            "evidence_path": f"{path}#chunk={chunk_id}",
            "preview": self._extract_preview(chunk_text),
            "reading_tags": dedup_tags,
        }

    # ------------------------
    # Public helpers
    # ------------------------
    def chunk_to_evidence(self, chunk: DocChunk) -> List[str]:
        if not isinstance(chunk, DocChunk):
            return []
        return [chunk.evidence_id]

    def chunk_to_source(self, chunk: DocChunk) -> str:
        if not isinstance(chunk, DocChunk):
            return "reading"
        return chunk.source_label

    def chunk_to_memory_stub(self, chunk: DocChunk) -> Dict[str, Any]:
        meta = dict(chunk.meta or {})
        return {
            "source": str(meta.get("source_label") or "reading"),
            "evidence": [str(meta.get("evidence_id") or f"{chunk.doc_id}:chunk_{chunk.chunk_id}")],
            "domain": meta.get("domain_hint"),
            "tags": list(meta.get("reading_tags", [])),
            "details": {
                "doc_id": chunk.doc_id,
                "chunk_index": chunk.chunk_id,
                "title": meta.get("title"),
                "file_name": meta.get("file_name"),
                "file_type": meta.get("file_type"),
                "path": meta.get("path"),
                "heading_hint": meta.get("heading_hint"),
                "chapter_hint": meta.get("chapter_hint"),
                "section_hint": meta.get("section_hint"),
                "book_title_hint": meta.get("book_title_hint"),
                "reader_name": meta.get("reader_name"),
                "preview": meta.get("preview"),
            },
        }

    def make_chunks(
        self,
        path: str,
        *,
        words_per_chunk: int = 900,
        overlap_words: int = 120,
    ) -> List[DocChunk]:
        path = self._clean_path(path)
        text = self.read_document(path)
        if not text:
            return []

        file_name = os.path.basename(path)
        file_stem, _ = os.path.splitext(file_name)
        doc_id = self._build_doc_id(path)

        pieces = self.chunk_text(
            text,
            words_per_chunk=words_per_chunk,
            overlap_words=overlap_words,
        )
        if not pieces:
            return []

        total_chunks = len(pieces)
        out: List[DocChunk] = []

        for i, t in enumerate(pieces):
            meta = self._make_chunk_meta(
                path=path,
                doc_id=doc_id,
                file_stem=file_stem,
                chunk_id=i,
                chunk_text=t,
                total_chunks=total_chunks,
            )
            out.append(DocChunk(doc_id=doc_id, chunk_id=i, text=t, meta=meta))

        return out

    def read_for_prompt(
        self,
        path: str,
        *,
        preview_chars: int = 1600,
        chunk_words: int = 700,
        max_chunks: int = 3,
    ) -> Dict[str, Any]:
        path = self._clean_path(path)
        if not path or not os.path.exists(path):
            return {
                "exists": False,
                "read_ok": False,
                "path": path,
                "file_type": None,
                "text": "",
                "chunk_count": 0,
                "chunk_previews": [],
                "error": "file does not exist",
            }

        text = self.read_document(path)
        if not text:
            return {
                "exists": True,
                "read_ok": False,
                "path": path,
                "file_type": self._infer_file_type_label(path),
                "text": "",
                "chunk_count": 0,
                "chunk_previews": [],
                "error": "no readable text extracted",
            }

        chunks = self.make_chunks(
            path,
            words_per_chunk=chunk_words,
            overlap_words=100,
        )

        chunk_previews: List[str] = []
        for chunk in chunks[:max_chunks]:
            label = chunk.meta.get("chunk_label", f"chunk_{chunk.chunk_id}")
            preview = self._extract_preview(chunk.text, max_chars=500)
            chunk_previews.append(f"[{label}] {preview}")

        return {
            "exists": True,
            "read_ok": True,
            "path": path,
            "file_type": self._infer_file_type_label(path),
            "text": self._extract_preview(text, max_chars=preview_chars),
            "chunk_count": len(chunks),
            "chunk_previews": chunk_previews,
            "error": None,
        }

    def select_relevant_chunks(
        self,
        chunks: List[DocChunk],
        query: str,
        *,
        top_k: int = 3,
    ) -> List[DocChunk]:
        q_tokens = set(self._TOKEN_RE.findall((query or "").lower()))
        if not q_tokens or not chunks:
            return []

        scored: List[tuple[int, DocChunk]] = []

        for chunk in chunks:
            c_tokens = set(self._TOKEN_RE.findall((chunk.text or "").lower()))
            overlap = len(q_tokens & c_tokens)
            scored.append((overlap, chunk))

        scored.sort(key=lambda x: x[0], reverse=True)
        return [chunk for score, chunk in scored[:top_k] if score > 0]